import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to Word table"""
    if not table_lines:
        return
    
    # Parse header
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
    
    # Parse rows
    rows_data = []
    for line in table_lines[2:]:  # Skip separator line
        if '|' in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if row:
                rows_data.append(row)
    
    if not headers or not rows_data:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add data rows
    for row_idx, row_data in enumerate(rows_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_data
    
    doc.add_paragraph()  # Add spacing after table

def parse_markdown_to_docx(markdown_file, docx_file):
    """Convert markdown file to DOCX"""
    
    if not os.path.exists(markdown_file):
        print(f"Error: File not found: {markdown_file}")
        return

    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse markdown content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Title (first # heading)
        if line.startswith('# ') and i == 0:
            title = line[2:].strip()
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.size = Pt(24)
                run.font.bold = True
            doc.add_paragraph()
            i += 1
            continue

        # Level 1 Heading (##)
        if line.startswith('## '):
            section = line[3:].strip()
            heading = doc.add_heading(section, level=1)
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            i += 1
            continue

        # Level 2 Heading (###) -> Map to Level 2 in Docx
        if line.startswith('### '):
            section = line[4:].strip()
            heading = doc.add_heading(section, level=2)
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
            continue
            
        # Fallback for other # usage
        if line.startswith('# '):
             section = line[2:].strip()
             heading = doc.add_heading(section, level=1)
             i += 1
             continue
        
        # Horizontal rule
        if line.startswith('---'):
            doc.add_paragraph()
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph(code_text)
                # Format as monospace
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            doc.add_paragraph()
            i += 1
            continue
        
        # Table detection
        if '|' in line and i + 1 < len(lines) and '---' in lines[i+1]:
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1
            
            add_table_from_markdown(doc, table_lines)
            i = j
            continue
        
        # Regular paragraph / List items
        if line:
            # Simple list handling
            if line.startswith('* ') or line.startswith('- '):
                # It's a list item
                line_content = line[2:]
                para = doc.add_paragraph(style='List Bullet')
                # Process formatting inside list item
                line = line_content
            elif re.match(r'^\d+\.\s', line):
                 # Numbered list
                 parts = line.split('.', 1)
                 line_content = parts[1].strip() if len(parts) > 1 else line
                 para = doc.add_paragraph(style='List Number')
                 line = line_content
            else:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Convert markdown formatting in the line
            line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line) # Bold
            # line = re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'<i>\1</i>', line) # Italic
            line = re.sub(r'`(.*?)`', r'<c>\1</c>', line) # Inline code
            
            parts = re.split(r'(<b>.*?</b>|<i>.*?</i>|<c>.*?</c>)', line)
            for part in parts:
                if not part:
                    continue
                if part.startswith('<b>') and part.endswith('</b>'):
                    text = part[3:-4]
                    run = para.add_run(text)
                    run.font.bold = True
                elif part.startswith('<i>') and part.endswith('</i>'):
                    text = part[3:-4]
                    run = para.add_run(text)
                    run.font.italic = True
                elif part.startswith('<c>') and part.endswith('</c>'):
                    text = part[3:-4]
                    run = para.add_run(text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(220, 50, 50) # Red-ish for code
                else:
                    part = part.replace('<br>', '\n')
                    clean_text = re.sub(r'<[^>]+>', '', part)
                    if clean_text:
                        para.add_run(clean_text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"DOCX created successfully: {docx_file}")

if __name__ == "__main__":
    base_dir = r"c:\Users\admin\Downloads\O2AI\Deploy\O2AI-Fax_Automation"
    markdown_file = os.path.join(base_dir, "EPIC_SIMPLE_SETUP.md")
    docx_file = os.path.join(base_dir, "EPIC_SIMPLE_SETUP.docx")
    parse_markdown_to_docx(markdown_file, docx_file)
